#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""GWMD → docx 转换器（党政机关公文格式 GB/T 9704-2012）

用法: python3 gwmd2docx.py <输入.gwmd> <输出.docx>

版面要求：
  - 页边距：上 37mm / 下 35mm / 左 28mm / 右 26mm（版心 156×225mm）
  - 每面 22 行，每行 28 个字（正文 3 号仿宋，行距固定值 29pt）
  - 脚注为 docx 页脚真注
"""
import re
import sys
import shutil
import zipfile
import unicodedata
from xml.sax.saxutils import escape
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ---------------------------------------------------------------- 常量
FONT_TITLE = '方正小标宋简体'      # 大标题 2号
FONT_H1 = '黑体'                   # 一级标题 3号
FONT_H2 = '楷体'                   # 二级标题 3号
FONT_BODY = '仿宋'                 # 正文 3号
FONT_PAGE = '宋体'                 # 页码 4号
FONT_ASCII = 'Times New Roman'     # 公文数字/字母（除页码外）
SIZE_TITLE = 22                    # 2号
SIZE_H1 = 16                       # 3号
SIZE_BODY = 16                     # 3号
SIZE_PAGE = 14                     # 4号
LINE_SPACING = 29                  # 正文行距（磅，保证每面22行）
TITLE_LINE_SPACING = 33            # 标题行距
RED = RGBColor(0xFF, 0x00, 0x00)

MARGIN_TOP = 3.7
MARGIN_BOTTOM = 3.5
MARGIN_LEFT = 2.8
MARGIN_RIGHT = 2.6
FOOTER_DIST = 2.5

# ---------------------------------------------------------------- 行内解析
INLINE_BOLD = re.compile(r'\*\*(.+?)\*\*')
INLINE_ITALIC = re.compile(r'(?<!\*)\*([^*]+?)\*(?!\*)')
INLINE_FOOTNOTE = re.compile(r'\[\^(\d+)\]')
INLINE_LINK = re.compile(r'\[([^\]]+)\]\((https?://[^)]+)\)')

# ---- 转义：\X 输出字面 X（特殊字符表顺序与前端 app.js ESC_CHARS 严格一致）----
ESC_CHARS = '\\*[]^@#|>('
ESC_PUA = [chr(0xE000 + i) for i in range(len(ESC_CHARS))]


def protect_escapes(s):
    """把 \\X（X 为特殊字符）替换为私有区占位，防止后续标记解析"""
    out = []
    i = 0
    n = len(s)
    while i < n:
        if s[i] == '\\' and i + 1 < n:
            idx = ESC_CHARS.find(s[i + 1])
            if idx >= 0:
                out.append(ESC_PUA[idx])
                i += 2
                continue
        out.append(s[i])
        i += 1
    return ''.join(out)


def restore_escapes(s):
    """把占位符恢复为字面字符"""
    for pua, ch in zip(ESC_PUA, ESC_CHARS):
        s = s.replace(pua, ch)
    return s


def split_footnotes(text):
    """把 [^n] 从文本中拆出，返回 [(seg, footnote_num_or_None)]"""
    parts = []
    pos = 0
    for m in INLINE_FOOTNOTE.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], None))
        parts.append(('', int(m.group(1))))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], None))
    return parts


def set_font(run, name, size, bold=False, color=None, ascii_font=FONT_ASCII):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    # 字符间距紧缩 -0.25pt（每字约15.75pt），保证每行28字（版心156mm≈442.2pt）
    spacing = rPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        rPr.append(spacing)
    spacing.set(qn('w:val'), '-5')


def set_spacing(para, pt=LINE_SPACING):
    pf = para.paragraph_format
    pf.line_spacing_rule = None
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:beforeLines'), '0')
    spacing.set(qn('w:afterLines'), '0')


def set_first_line_chars(para, chars=200):
    """首行缩进 N 字符（1/100 字符单位）"""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars))
    ind.set(qn('w:firstLine'), str(int(chars / 100 * SIZE_BODY * 20)))


def set_right_indent_chars(para, width_100):
    """右缩进（1/100 字符宽：半角空格=50，全角空格=100；右对齐落款整体左移）"""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:rightChars'), str(width_100))
    ind.set(qn('w:right'), str(int(width_100 / 100 * SIZE_BODY * 20)))


def add_footnote_ref(para, fid):
    """在段落中插入脚注引用 run（三号上标数字）"""
    run = para.add_run()
    r = run._element
    rPr = OxmlElement('w:rPr')
    va = OxmlElement('w:vertAlign'); va.set(qn('w:val'), 'superscript')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '32')
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '32')
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:ascii'), FONT_ASCII); rF.set(qn('w:hAnsi'), FONT_ASCII); rF.set(qn('w:eastAsia'), FONT_BODY)
    rPr.append(va); rPr.append(sz); rPr.append(szCs); rPr.append(rF)
    r.append(rPr)
    ref = OxmlElement('w:footnoteReference')
    ref.set(qn('w:id'), str(fid))
    r.append(ref)


def add_text_runs(para, text, font=FONT_BODY, size=SIZE_BODY, bold_all=False,
                    ascii_font=FONT_ASCII):
    """添加带行内标记（加粗/斜体/脚注/链接）的 run，链接为真实超链接；\\X 为转义"""
    text = protect_escapes(text)
    for seg, fn in split_footnotes(text):
        if fn is not None:
            add_footnote_ref(para, fn)
            continue
        if not seg:
            continue
        # 拆出链接段 [文字](url)，其余按普通文本处理
        pos = 0
        for m in INLINE_LINK.finditer(seg):
            if m.start() > pos:
                add_plain_runs(para, seg[pos:m.start()], font, size, bold_all, ascii_font)
            add_hyperlink_runs(para, m.group(1), m.group(2), font, size, bold_all, ascii_font)
            pos = m.end()
        if pos < len(seg):
            add_plain_runs(para, seg[pos:], font, size, bold_all, ascii_font)


def add_plain_runs(para, seg, font=FONT_BODY, size=SIZE_BODY, bold_all=False,
                     ascii_font=FONT_ASCII):
    """普通文本 run（加粗/斜体标记）"""
    tokens = INLINE_BOLD.split(seg)
    for idx, tok in enumerate(tokens):
        if not tok:
            continue
        is_bold = (idx % 2 == 1)
        for j, sub in enumerate(INLINE_ITALIC.split(tok)):
            if not sub:
                continue
            run = para.add_run(restore_escapes(sub))
            set_font(run, font, size, bold=bold_all or is_bold, ascii_font=ascii_font)
            if j % 2 == 1:
                run.font.italic = True


def add_hyperlink_runs(para, text, url, font=FONT_BODY, size=SIZE_BODY, bold_all=False,
                          ascii_font=FONT_ASCII):
    """超链接 run（w:hyperlink + External 关系，保留字体/加粗/斜体/紧缩）"""
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    tokens = INLINE_BOLD.split(text)
    for idx, tok in enumerate(tokens):
        if not tok:
            continue
        is_bold = (idx % 2 == 1)
        for j, sub in enumerate(INLINE_ITALIC.split(tok)):
            if not sub:
                continue
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), font)
            rFonts.set(qn('w:ascii'), ascii_font)
            rFonts.set(qn('w:hAnsi'), ascii_font)
            rPr.append(rFonts)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
            szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(int(size * 2))); rPr.append(szCs)
            spacing = OxmlElement('w:spacing'); spacing.set(qn('w:val'), '-5'); rPr.append(spacing)
            if bold_all or is_bold:
                rPr.append(OxmlElement('w:b'))
            if j % 2 == 1:
                rPr.append(OxmlElement('w:i'))
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = restore_escapes(sub)
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            hyperlink.append(r)
    para._p.append(hyperlink)


# ---------------------------------------------------------------- 块解析
RE_DIRECTIVE = re.compile(r'^@([a-z]+)\s*(.*)$', re.S)
RE_HEADING = re.compile(r'^(#{1,6})\s+(.*)$')
RE_ORDERED = re.compile(r'^(\d+)[\.、]\s+(.*)$')
RE_UNORDERED = re.compile(r'^[-*]\s+(.*)$')
RE_QUOTE = re.compile(r'^>\s?(.*)$')
RE_TABLE_ROW = re.compile(r'^\|.*\|$')
RE_TABLE_SEP = re.compile(r'^\|[\s:\-|]+\|$')
RE_FOOTNOTE_DEF = re.compile(r'^\[\^(\d+)\]:\s*(.*)$')

HEADING_MAP = {1: 'h1', 2: 'h2', 3: 'h3', 4: 'h4', 5: 'h4', 6: 'h4'}
HEADER_DIRECTIVES = {
    'header': 'masthead', 'docno': 'docno', 'secret': 'secret',
    'urgent': 'urgent', 'serial': 'serial', 'signer': 'signer',
}
BODY_DIRECTIVES = {
    'to': 'to', 'sign': 'sign', 'date': 'date',
    'stamp': 'stamp', 'note': 'note', 'subtitle': 'subtitle',
    'title': 'title',
}
RECORD_DIRECTIVES = {'cc': 'cc', 'print': 'print', 'issue': 'issue'}

# 附件名称提取：@attach [附件1] [附件2] ...
RE_ATTACH_ITEM = re.compile(r'\[([^\]]+)\]')

# 配对标题：# 标题 # 正文
RE_INLINE_TITLE = re.compile(r'^(#{1,5})\s+(.+?)\s+\1\s+(.+)$')

# ---- 标题自动编号 ----
CN_DIGITS = '零一二三四五六七八九'
RE_CN_NUM = re.compile(r'^([一二三四五六七八九十]+)、')
RE_CN_PAREN = re.compile(r'^（([一二三四五六七八九十]+)）')
RE_ARABIC_NUM = re.compile(r'^(\d+)[\.．]')
RE_ARABIC_PAREN = re.compile(r'^（(\d+)）')


def int_to_cn(n):
    """整数 → 中文数字（支持 1~99）"""
    if n <= 0:
        return '零'
    if n < 10:
        return CN_DIGITS[n]
    if n < 20:
        return '十' + (CN_DIGITS[n - 10] if n > 10 else '')
    tens, ones = divmod(n, 10)
    return CN_DIGITS[tens] + '十' + (CN_DIGITS[ones] if ones else '')


def cn_to_int(s):
    """中文数字 → 整数（支持 1~99）"""
    if not s:
        return 0
    units = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
             '六': 6, '七': 7, '八': 8, '九': 9}
    if '十' not in s:
        return units.get(s, 0)
    parts = s.split('十')
    tens = units.get(parts[0], 1) if parts[0] else 1
    ones = units.get(parts[1], 0) if len(parts) > 1 and parts[1] else 0
    return tens * 10 + ones


def num_prefix(level, n):
    if level == 2:
        return int_to_cn(n) + '、'
    if level == 3:
        return '（' + int_to_cn(n) + '）'
    if level == 4:
        return f'{n}.'
    if level == 5:
        return f'（{n}）'
    return ''


def detect_manual_number(text, level):
    """检测标题是否已手写编号，返回 (是否手写, 数值)"""
    if level == 2:
        m = RE_CN_NUM.match(text)
        if m:
            return True, cn_to_int(m.group(1))
    elif level == 3:
        m = RE_CN_PAREN.match(text)
        if m:
            return True, cn_to_int(m.group(1))
    elif level == 4:
        m = RE_ARABIC_NUM.match(text)
        if m:
            return True, int(m.group(1))
    elif level == 5:
        m = RE_ARABIC_PAREN.match(text)
        if m:
            return True, int(m.group(1))
    return False, 0


def number_headings(blocks):
    """标题自动编号：一、/（一）/1./（1），低层级随高层级重置"""
    counters = {2: 0, 3: 0, 4: 0, 5: 0}
    for b in blocks:
        t = b.get('type')
        if t in ('h1', 'h2', 'h3', 'h4'):
            level = {'h1': 2, 'h2': 3, 'h3': 4, 'h4': 5}[t]
        elif t == 'inline_title':
            level = b.get('level', 2)
        else:
            continue
        if t == 'inline_title':
            text = b.get('title', '').strip()
        else:
            text = b.get('text', '').strip()
        manual, num = detect_manual_number(text, level)
        if manual:
            counters[level] = max(counters[level], num)
        else:
            counters[level] += 1
            if t == 'inline_title':
                b['title'] = num_prefix(level, counters[level]) + text
            else:
                b['text'] = num_prefix(level, counters[level]) + text
        for lv in range(level + 1, 6):
            counters[lv] = 0
    return blocks


def parse_gwmd(text):
    """解析 GWMD 文本 → 块列表。块: {type, text, meta}"""
    blocks = []
    lines = text.split('\n')
    i = 0
    n = len(lines)
    footnotes = {}
    pending_table = None
    pending_list = None

    def flush_table():
        nonlocal pending_table
        if pending_table:
            blocks.append({'type': 'table', 'rows': pending_table})
            pending_table = None

    def flush_list():
        nonlocal pending_list
        if pending_list:
            blocks.append({'type': 'list', 'ordered': pending_list[0],
                           'items': pending_list[1]})
            pending_list = None

    while i < n:
        raw = lines[i]
        line = raw.strip()
        i += 1

        if not line:
            flush_table(); flush_list()
            continue

        # 行级转义：\X 开头 → 字面正文（去掉反斜杠）
        if line.startswith('\\'):
            flush_table(); flush_list()
            blocks.append({'type': 'body', 'text': line[1:]})
            continue

        # 注释行：<!-- 内容 --> 不导出
        if line.startswith('<!--'):
            continue

        m = RE_FOOTNOTE_DEF.match(line)
        if m:
            footnotes[m.group(1)] = m.group(2).strip()
            continue

        if line.startswith('@'):
            flush_table(); flush_list()
            # 用原始行匹配，保留行末空格（@sign/@date 手动对齐用）
            raw_line = raw.rstrip('\r\n')
            m = RE_DIRECTIVE.match(raw_line.lstrip())
            if m:
                key = m.group(1)
                if key in ('sign', 'date'):
                    content = m.group(2)  # 保留行末空格
                else:
                    content = m.group(2).strip()
                if key == 'comment':
                    pass  # 注释，不导出
                elif key == 'redline':
                    blocks.append({'type': 'redline'})
                elif key == 'pagebreak':
                    blocks.append({'type': 'pagebreak'})
                elif key == 'config':
                    blocks.append({'type': 'config', 'text': content})
                elif key in ('blank', 'spacer', 'space', 'sp'):
                    cnt = 1
                    if content:
                        try:
                            cnt = int(content)
                        except ValueError:
                            cnt = 1
                    blocks.append({'type': 'blank', 'count': max(1, min(cnt, 20))})
                elif key == 'attach':
                    items = RE_ATTACH_ITEM.findall(content)
                    if items:
                        blocks.append({'type': 'attach', 'items': [i.strip() for i in items]})
                    else:
                        name = re.sub(r'^附件\s*[:：]?\s*', '', content).strip()
                        blocks.append({'type': 'attach', 'items': [name] if name else ['附件']})
                elif key in HEADER_DIRECTIVES:
                    blocks.append({'type': HEADER_DIRECTIVES[key], 'text': content})
                elif key in BODY_DIRECTIVES:
                    blocks.append({'type': BODY_DIRECTIVES[key], 'text': content})
                elif key in RECORD_DIRECTIVES:
                    blocks.append({'type': RECORD_DIRECTIVES[key], 'text': content})
                else:
                    blocks.append({'type': 'body', 'text': line})
            continue

        if RE_TABLE_ROW.match(line):
            flush_list()
            if RE_TABLE_SEP.match(line):
                continue
            cells = [c.strip() for c in protect_escapes(line.strip('|')).split('|')]
            if pending_table is None:
                pending_table = []
            pending_table.append(cells)
            continue

        # 配对标题：### 标题 ### 正文（标题与正文同一段）
        m = RE_INLINE_TITLE.match(line)
        if m:
            flush_table(); flush_list()
            level = min(len(m.group(1)) + 1, 5)
            blocks.append({'type': 'inline_title', 'level': level,
                           'title': m.group(2).strip(), 'body': m.group(3).strip()})
            continue

        m = RE_HEADING.match(line)
        if m:
            flush_table(); flush_list()
            level = len(m.group(1))
            blocks.append({'type': HEADING_MAP.get(level, 'h4'),
                           'text': m.group(2).strip()})
            continue

        m = RE_ORDERED.match(line)
        if m:
            flush_table()
            if pending_list is None:
                pending_list = [True, []]
            elif not pending_list[0]:
                flush_list(); pending_list = [True, []]
            pending_list[1].append(m.group(2))
            continue
        m = RE_UNORDERED.match(line)
        if m:
            flush_table()
            if pending_list is None:
                pending_list = [False, []]
            elif pending_list[0]:
                flush_list(); pending_list = [False, []]
            pending_list[1].append(m.group(1))
            continue

        m = RE_QUOTE.match(line)
        if m:
            flush_table(); flush_list()
            blocks.append({'type': 'quote', 'text': m.group(1)})
            continue

        flush_table(); flush_list()
        para_lines = [line]
        while i < n:
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith('#') or nxt.startswith('@') or nxt.startswith('|') or \
               nxt.startswith('>') or is_list_start(nxt) or \
               RE_FOOTNOTE_DEF.match(nxt) or RE_TABLE_ROW.match(nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append({'type': 'body', 'text': ' '.join(para_lines)})

    flush_table(); flush_list()
    blocks = strip_undefined_footnotes(number_headings(blocks), footnotes)
    blocks = extract_trailing_indent(blocks)
    return blocks, footnotes


def strip_undefined_footnotes(blocks, footnotes):
    """剔除正文中未定义（无 [^n]: 脚注定义）的引用，防止 docx 脚注损坏"""
    def clean(s):
        if not s:
            return s
        return INLINE_FOOTNOTE.sub(lambda m: m.group(0) if m.group(1) in footnotes else '', s)
    for b in blocks:
        t = b.get('type')
        if t == 'inline_title':
            b['title'] = clean(b.get('title', ''))
            b['body'] = clean(b.get('body', ''))
        elif t == 'table':
            for row in b.get('rows', []):
                for i, c in enumerate(row):
                    row[i] = clean(c)
        elif t == 'list':
            b['items'] = [clean(i) for i in b.get('items', [])]
        else:
            b['text'] = clean(b.get('text', ''))
    return blocks


def extract_trailing_indent(blocks):
    """@sign/@date 行末空格 → right 字段（1/100 字符宽：半角空格=50，全角空格=100）"""
    for b in blocks:
        if b.get('type') not in ('sign', 'date'):
            continue
        t = b.get('text', '')
        m = re.search(r'[ \u3000]+$', t)
        if m:
            trail = m.group(0)
            b['right'] = trail.count(' ') * 50 + trail.count('\u3000') * 100
            b['text'] = t[:m.start()]
    return blocks


def is_list_start(line):
    return bool(RE_ORDERED.match(line)) or bool(RE_UNORDERED.match(line))


# ---- 结构校验（不阻断导出；与前端 app.js validate 逻辑严格一致）----
KNOWN_DIRECTIVES = (set(HEADER_DIRECTIVES) | set(BODY_DIRECTIVES) | set(RECORD_DIRECTIVES)
                    | {'redline', 'pagebreak', 'config', 'blank', 'spacer', 'space',
                       'sp', 'attach', 'comment'})


def validate_gwmd(text):
    """结构校验：返回警告列表（用于预览告警条与一致性测试）"""
    warnings = []
    blocks, footnotes = parse_gwmd(text)

    # 大标题
    titles = [b for b in blocks if b.get('type') == 'title']
    if not titles:
        warnings.append('缺少公文大标题（@title）')
    elif len(titles) > 1:
        warnings.append('检测到多个 @title，公文大标题应只有一个')

    # 未定义脚注（原始文本扫描，排除转义 \[^n\]）
    defined = set(footnotes)
    used = set(re.findall(r'(?<!\\)\[\^(\d+)\]', text))
    for n in sorted(used - defined, key=int):
        warnings.append(f'脚注 [^{n}] 无定义，导出时已忽略')

    # 未知指令（行级扫描）
    for raw_line in text.splitlines():
        l = raw_line.lstrip()
        if l.startswith('@'):
            m = RE_DIRECTIVE.match(l)
            if m and m.group(1) not in KNOWN_DIRECTIVES:
                warnings.append(f'未知指令 @{m.group(1)}，已按正文处理')

    # 标题层级跳跃
    prev = 1
    for b in blocks:
        t = b.get('type')
        if t in ('h1', 'h2', 'h3', 'h4'):
            lv = {'h1': 2, 'h2': 3, 'h3': 4, 'h4': 5}[t]
            title = b.get('text', '')
        elif t == 'inline_title':
            lv = b.get('level', 2)
            title = b.get('title', '')
        else:
            continue
        if lv > prev + 1:
            warnings.append(f'标题层级跳跃："{title}"，前面缺少上级标题')
        prev = max(prev, lv)

    # @stamp 之后仍有正文
    stamp_seen = False
    for b in blocks:
        t = b.get('type')
        if t == 'stamp':
            stamp_seen = True
            continue
        if stamp_seen and t in ('body', 'title', 'h1', 'h2', 'h3', 'h4',
                                'inline_title', 'to', 'subtitle', 'attach'):
            warnings.append('@stamp 之后仍有正文内容，印章应位于落款之后、附注之前')
            break

    return warnings


# ---------------------------------------------------------------- 渲染
def render_block(doc, block):
    btype = block.get('type')
    text = block.get('text', '')

    if btype == 'pagebreak':
        doc.add_page_break()
        return
    if btype == 'blank':
        for _ in range(block.get('count', 1)):
            p = doc.add_paragraph()
            set_spacing(p)
        return
    if btype == 'redline':
        # 红线：极矮段落（2pt）+ 段落下边框，紧贴签发人下方，不留空行
        p = doc.add_paragraph()
        set_spacing(p, 2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '18')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'FF0000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return
    if btype == 'table':
        rows = block['rows']
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                cell = table.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_spacing(p, 20)
                add_text_runs(p, row[ci] if ci < len(row) else '', font=FONT_BODY, size=12, bold_all=(ri == 0))
        doc.add_paragraph()
        return
    if btype == 'list':
        ordered, items = block['ordered'], block['items']
        for idx, item in enumerate(items, 1):
            p = doc.add_paragraph()
            set_spacing(p)
            set_first_line_chars(p, 200)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            prefix = f'{idx}. ' if ordered else '- '
            add_text_runs(p, prefix + item, font=FONT_BODY, size=SIZE_BODY)
        return
    if btype == 'quote':
        p = doc.add_paragraph()
        set_spacing(p)
        set_first_line_chars(p, 200)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_text_runs(p, text, font=FONT_H2, size=SIZE_BODY)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '8')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '000000')
        pBdr.append(left)
        pPr.append(pBdr)
        return

    p = doc.add_paragraph()

    if btype == 'title':
        # 大标题支持 | 手工换行：同一段落内 <w:br/> 换行，保持居中
        set_spacing(p, TITLE_LINE_SPACING)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_parts = protect_escapes(text).split('|')
        for ti, tpart in enumerate(title_parts):
            if ti > 0:
                br_run = p.add_run()
                set_font(br_run, FONT_TITLE, SIZE_TITLE, bold=True)
                br_run.add_break()
            if tpart:
                add_text_runs(p, tpart, font=FONT_TITLE, size=SIZE_TITLE, bold_all=True)
    elif btype == 'subtitle':
        set_spacing(p, TITLE_LINE_SPACING)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_runs(p, text, font=FONT_H2, size=SIZE_BODY)
    elif btype == 'h1':
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_first_line_chars(p, 200)
        add_text_runs(p, text, font=FONT_H1, size=SIZE_H1, bold_all=True)
    elif btype == 'h2':
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_first_line_chars(p, 200)
        add_text_runs(p, text, font=FONT_H2, size=SIZE_BODY)
    elif btype in ('h3', 'h4'):
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_first_line_chars(p, 200)
        add_text_runs(p, text, font=FONT_BODY, size=SIZE_BODY, bold_all=(btype == 'h3'))
    elif btype == 'inline_title':
        # 配对标题：# 标题 # 正文（标题与正文同一段）
        level = block.get('level', 2)
        title = block.get('title', '').strip()
        body = block.get('body', '').strip()
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_first_line_chars(p, 200)
        if level == 2:
            add_text_runs(p, title, font=FONT_H1, size=SIZE_H1, bold_all=True)
        elif level == 3:
            add_text_runs(p, title, font=FONT_H2, size=SIZE_BODY)
        elif level == 4:
            add_text_runs(p, title, font=FONT_BODY, size=SIZE_BODY, bold_all=True)
        else:
            add_text_runs(p, title, font=FONT_BODY, size=SIZE_BODY)
        if body:
            add_text_runs(p, body, font=FONT_BODY, size=SIZE_BODY)
    elif btype == 'to':
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_text_runs(p, text, font=FONT_BODY, size=SIZE_BODY)
    elif btype in ('sign', 'date'):
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if block.get('right'):
            set_right_indent_chars(p, block['right'])
        add_text_runs(p, text, font=FONT_BODY, size=SIZE_BODY)
    elif btype == 'attach':
        # 附件说明：与网页预览一致：首行缩进2字符；多附件后续行缩进5字符
        items = block.get('items') or ([text] if text else ['附件'])
        for i, name in enumerate(items):
            p = doc.add_paragraph()
            set_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_first_line_chars(p, 200 if i == 0 else 500)
            label = f'附件：{i + 1}. ' if i == 0 else f'{i + 1}. '
            add_text_runs(p, label + name, font=FONT_BODY, size=SIZE_BODY)
    elif btype == 'note':
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_first_line_chars(p, 200)
        add_text_runs(p, text, font=FONT_BODY, size=SIZE_BODY)
    elif btype == 'masthead':
        set_spacing(p, TITLE_LINE_SPACING)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_runs(p, text, font=FONT_TITLE, size=SIZE_TITLE, bold_all=True)
        for r in p.runs:
            set_font(r, FONT_TITLE, SIZE_TITLE, bold=True, color=RED)
    elif btype == 'docno':
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_runs(p, text, font=FONT_BODY, size=SIZE_BODY)
    elif btype in ('secret', 'urgent'):
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_text_runs(p, text, font=FONT_H1, size=SIZE_H1)
    elif btype == 'serial':
        # 份号数字用黑体（GB/T 9704：份号为 3 号黑体阿拉伯数字，不走 TNR）
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_text_runs(p, text, font=FONT_H1, size=SIZE_H1, ascii_font=FONT_H1)
    elif btype == 'signer':
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text_runs(p, text, font=FONT_BODY, size=SIZE_BODY)
    elif btype == 'stamp':
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text_runs(p, '（印章）', font=FONT_BODY, size=SIZE_BODY)
    elif btype in ('cc', 'print', 'issue'):
        set_spacing(p, 24)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 版记 12pt，与预览端 CSS（#preview .gw-cc/.gw-print/.gw-issue）严格一致。
        # 关闭中西文自动间距（autoSpaceDE/DN）与字符网格对齐（snapToGrid）：
        # Word 默认在中英/中文与数字交界处自动加约 1/8 字宽间隙，多个交界累计
        # 接近一个汉字宽，导致同一内容 docx 比预览早折行一个字。
        add_text_runs(p, text, font=FONT_BODY, size=12)
        pPr = p._p.get_or_add_pPr()
        for _tag in ('w:autoSpaceDE', 'w:autoSpaceDN'):
            _el = OxmlElement(_tag)
            _el.set(qn('w:val'), '0')
            pPr.insert_element_before(_el, 'w:bidi', 'w:adjustRightInd',
                                      'w:snapToGrid', 'w:spacing', 'w:ind',
                                      'w:contextualSpacing', 'w:jc')
        _el = OxmlElement('w:snapToGrid')
        _el.set(qn('w:val'), '0')
        pPr.insert_element_before(_el, 'w:spacing', 'w:ind',
                                  'w:contextualSpacing', 'w:jc')
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), '000000')
        pBdr.append(top)
        pPr.append(pBdr)
    elif btype == 'config':
        return
    else:  # body
        set_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_first_line_chars(p, 200)
        add_text_runs(p, text, font=FONT_BODY, size=SIZE_BODY)


def add_page_numbers(doc):
    """公文页码：4号宋体，— N —，单页码居右空一字，双页码居左空一字"""
    try:
        doc.settings.odd_and_even_pages_header_footer = True
        for section in doc.sections:
            for footer, is_even in ((section.footer, False),
                                    (section.even_page_footer, True)):
                footer.is_linked_to_previous = False
                for p in list(footer.paragraphs):
                    p._p.getparent().remove(p._p)
                para = footer.add_paragraph()
                pf = para.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                para.alignment = (WD_ALIGN_PARAGRAPH.LEFT if is_even
                                  else WD_ALIGN_PARAGRAPH.RIGHT)
                pPr = para._p.get_or_add_pPr()
                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                if is_even:
                    ind.set(qn('w:leftChars'), '100')
                else:
                    ind.set(qn('w:rightChars'), '100')

                def style_run(run):
                    run.font.size = Pt(SIZE_PAGE)
                    run.font.name = FONT_PAGE
                    rPr = run._element.get_or_add_rPr()
                    rF = rPr.find(qn('w:rFonts'))
                    if rF is None:
                        rF = OxmlElement('w:rFonts')
                        rPr.append(rF)
                    rF.set(qn('w:eastAsia'), FONT_PAGE)
                    rF.set(qn('w:ascii'), FONT_PAGE)
                    rF.set(qn('w:hAnsi'), FONT_PAGE)

                r1 = para.add_run('— ')
                style_run(r1)
                r2 = para.add_run()
                fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
                ite = OxmlElement('w:instrText'); ite.set(qn('xml:space'), 'preserve'); ite.text = ' PAGE '
                fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
                r2._r.append(fc1); r2._r.append(ite); r2._r.append(fc2)
                style_run(r2)
                r3 = para.add_run(' —')
                style_run(r3)
    except Exception as e:
        print(f'警告：页码生成失败：{e}', file=sys.stderr)


# ---------------------------------------------------------------- 真脚注注入
def inject_footnotes(docx_path, footnotes):
    """把 footnotes.xml 注入 docx，实现页面底部真脚注"""
    if not footnotes:
        return
    tmp = docx_path + '.fntmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        items = zin.namelist()
        data = {name: zin.read(name) for name in items}

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    parts = []
    parts.append(
        f'<w:footnote w:type="separator" w:id="-1">'
        f'<w:p><w:pPr><w:spacing w:after="0" w:line="400" w:lineRule="exact"/></w:pPr>'
        f'<w:r><w:separator/></w:r></w:p></w:footnote>'
    )
    parts.append(
        f'<w:footnote w:type="continuationSeparator" w:id="0">'
        f'<w:p><w:pPr><w:spacing w:after="0" w:line="400" w:lineRule="exact"/></w:pPr>'
        f'<w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
    )
    for fid in sorted(footnotes, key=int):
        t = escape(footnotes[fid])
        parts.append(
            f'<w:footnote w:id="{fid}">'
            f'<w:p><w:pPr>'
            f'<w:spacing w:after="0" w:line="400" w:lineRule="exact"/>'
            f'<w:ind w:left="0" w:hanging="0"/>'
            f'<w:jc w:val="left"/>'
            f'</w:pPr>'
            f'<w:r><w:rPr>'
            f'<w:rFonts w:ascii="{FONT_ASCII}" w:hAnsi="{FONT_ASCII}" w:eastAsia="{FONT_BODY}"/>'
            f'<w:sz w:val="28"/><w:szCs w:val="28"/>'
            f'</w:rPr>'
            f'<w:t xml:space="preserve">{fid}. {t}</w:t>'
            f'</w:r>'
            f'</w:p></w:footnote>'
        )
    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:footnotes xmlns:w="{W}" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        + ''.join(parts) + '</w:footnotes>'
    )
    data['word/footnotes.xml'] = footnotes_xml.encode('utf-8')

    ct = data['[Content_Types].xml'].decode('utf-8')
    if 'footnotes.xml' not in ct:
        override = ('<Override PartName="/word/footnotes.xml" '
                    'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>')
        ct = ct.replace('</Types>', override + '</Types>')
        data['[Content_Types].xml'] = ct.encode('utf-8')

    rels = data['word/_rels/document.xml.rels'].decode('utf-8')
    if 'footnotes' not in rels:
        rel = ('<Relationship Id="rIdGWMDFootnotes" '
               'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
               'Target="footnotes.xml"/>')
        rels = rels.replace('</Relationships>', rel + '</Relationships>')
        data['word/_rels/document.xml.rels'] = rels.encode('utf-8')

    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, content in data.items():
            zout.writestr(name, content)
    shutil.move(tmp, docx_path)


# ---------------------------------------------------------------- 转换入口
def convert(gwmd_text, out_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(MARGIN_TOP)
    sec.bottom_margin = Cm(MARGIN_BOTTOM)
    sec.left_margin = Cm(MARGIN_LEFT)
    sec.right_margin = Cm(MARGIN_RIGHT)
    sec.footer_distance = Cm(FOOTER_DIST)

    style = doc.styles['Normal']
    style.font.name = FONT_ASCII
    style.font.size = Pt(SIZE_BODY)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    blocks, footnotes = parse_gwmd(gwmd_text)

    # 落款/日期对齐
    sign_texts = [b.get('text', '') for b in blocks if b.get('type') == 'sign']
    date_texts = [b.get('text', '') for b in blocks if b.get('type') == 'date']

    def dw(s):
        return sum(1.0 if unicodedata.east_asian_width(c) in ('F', 'W', 'A') else 0.5 for c in s)

    if sign_texts and date_texts:
        max_w = max([dw(t) for t in sign_texts] + [dw(t) for t in date_texts])
        sign_pad = {t: '　' * round(max_w - dw(t)) for t in sign_texts if dw(t) < max_w}
        date_pad = {t: '　' * round(max_w - dw(t)) for t in date_texts if dw(t) < max_w}
        for b in blocks:
            t = b.get('text', '')
            if b.get('type') == 'sign' and t in sign_pad:
                b['text'] = sign_pad[t] + t
            elif b.get('type') == 'date' and t in date_pad:
                b['text'] = date_pad[t] + t

    for block in blocks:
        render_block(doc, block)

    add_page_numbers(doc)
    doc.save(out_path)
    inject_footnotes(out_path, footnotes)
    return True


def main():
    if len(sys.argv) < 3:
        print('用法: python3 gwmd2docx.py <输入.gwmd> <输出.docx>', file=sys.stderr)
        return 1
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        text = f.read()
    convert(text, sys.argv[2])
    print(sys.argv[2])
    return 0


if __name__ == '__main__':
    sys.exit(main())
